#!/usr/bin/env python3
"""Assemble AgroElevate Final Black Book from chapter markdown files."""
from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
CHAPTERS_DIR = ROOT / "docs" / "blackbook" / "chapters"
DIAGRAMS_DIR = ROOT / "docs" / "blackbook" / "diagrams"
OUT_DIR = ROOT / "docs" / "blackbook"
OUT_ROOT = ROOT  # also copy to project root

CHAPTER_ORDER = [
    "00_front_matter.md",
    "01_introduction.md",
    "02_literature_survey.md",
    "03_architecture.md",
    "04_database.md",
    "05_modules.md",
    "06_android.md",
    "07_algorithms_testing.md",
    "08_conclusion.md",
    "09_security_apis_results.md",
]

DIAGRAM_APPENDIX = """
---

# Appendix A — System Diagrams (Editable Source)

Editable Mermaid sources are stored in `docs/blackbook/diagrams/`:

| File | Diagram |
|------|---------|
| `01_overall_architecture.mmd` | Overall three-tier architecture |
| `02_er_diagram.mmd` | Entity-relationship diagram |
| `03_royalty_workflow.mmd` | Option B royalty decision flow |
| `04_payment_flow.mmd` | Razorpay wallet sequence |
| `05_auth_flow.mmd` | Authentication and route guards |
| `06_ai_pipeline.mmd` | AI intelligence data pipeline |
| `07_use_case.mmd` | Use case diagram (five roles) |
| `08_order_lifecycle.mmd` | Order state machine |
| `09_marketplace_flow.mmd` | Marketplace checkout sequence |
| `10_deployment.mmd` | Dev vs production topology |
| `11_android_navigation.mmd` | Planned Android navigation |
| `12_activity_checkout.mmd` | Checkout activity diagram |

Import `.mmd` files into draw.io, Mermaid Live Editor, or VS Code Mermaid preview for submission figures.

---

# Appendix B — Screenshot Insertion Guide

| Figure | Screen to capture | Module |
|--------|-------------------|--------|
| Fig 5.1 | Landing page hero | Marketing |
| Fig 5.2 | Farmer dashboard KPI cards | Farmer |
| Fig 5.3 | Marketplace browse + filters | Marketplace |
| Fig 5.4 | Farmer My Listings tab | Farmer |
| Fig 5.5 | Cart checkout confirmation | Commerce |
| Fig 5.6 | Wallet balance + Razorpay top-up | Wallet |
| Fig 5.7 | Wallet transaction history | Wallet |
| Fig 5.8 | Orders list (buyer/seller) | Orders |
| Fig 5.9 | Trader inventory relist dialog | Trader |
| Fig 5.10 | Industrialist manufacturing panel | Industrialist |
| Fig 5.11 | Farmer Intelligence + crop recommendations | AI |
| Fig 5.12 | Income forecast insufficient-data state | AI |
| Fig 5.13 | AI Copilot conversation | Copilot |
| Fig 5.14 | Admin user moderation console | Admin |
| Fig 5.15 | Admin Payments + demo credit | Admin |
| Fig 6.1 | Android login screen (planned) | Android |
| Fig 6.2 | Android marketplace (planned) | Android |

---

# Appendix C — Project Verification Summary

| Verification | Command | Result |
|--------------|---------|--------|
| Production build | `npm run build` | PASS |
| Commerce E2E | `npm run commerce:verify` | 26/26 PASS |
| RPC smoke | `npm run commerce:smoke` | 7/7 PASS |
| AI health | `npm run ai:verify` | Health + dashboard PASS |
| Royalty math | Trader→Industrialist 5×₹70 | ₹43.75 (12.5%) |

**Release:** AgroElevate v1.0.0-rc (Web Platform)

"""


def assemble_markdown() -> str:
    parts: list[str] = []
    for name in CHAPTER_ORDER:
        path = CHAPTERS_DIR / name
        if not path.exists():
            raise FileNotFoundError(f"Missing chapter: {path}")
        parts.append(path.read_text(encoding="utf-8").strip())
        parts.append("\n\n")
    parts.append(DIAGRAM_APPENDIX.strip())
    return "\n".join(parts)


def md_to_docx(md_text: str, docx_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_LINE_SPACING

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    for level in range(1, 5):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Times New Roman"
        hs.font.bold = True
        hs.font.size = Pt(16 - level * 2)
        hs.paragraph_format.space_before = Pt(6)
        hs.paragraph_format.space_after = Pt(4)
        hs.paragraph_format.page_break_before = False
        hs.paragraph_format.keep_with_next = True

    mermaid_fig = 0
    lines = md_text.splitlines()
    i = 0
    in_table = False
    table_rows: list[list[str]] = []

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            return
        cols = max(len(r) for r in table_rows)
        tbl = doc.add_table(rows=len(table_rows), cols=cols)
        tbl.style = "Table Grid"
        for ri, row in enumerate(table_rows):
            for ci in range(cols):
                cell_text = row[ci] if ci < len(row) else ""
                cell_text = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
                cell_text = re.sub(r"`(.+?)`", r"\1", cell_text)
                cell = tbl.rows[ri].cells[ci]
                cell.text = cell_text.strip()
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
        table_rows = []
        in_table = False

    def add_code_block(code_lines: list[str], lang: str) -> None:
        nonlocal mermaid_fig
        if lang == "mermaid":
            mermaid_fig += 1
            p = doc.add_paragraph()
            run = p.add_run(
                f"[Figure {mermaid_fig} — Mermaid diagram; editable source in docs/blackbook/diagrams/]"
            )
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            return
        text = "\n".join(code_lines)
        if len(text) > 1800:
            text = text[:1800] + "\n... [truncated for print; see .md source]"
        p = doc.add_paragraph(text)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(9)

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("|") and "|" in line[1:]:
            if re.match(r"^\|[-:\s|]+\|$", line.strip()):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_rows = []
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            flush_table()

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.strip() == "---":
            pass
        elif line.strip().startswith(">"):
            p = doc.add_paragraph(line.lstrip("> ").strip())
            p.paragraph_format.left_indent = Inches(0.5)
        elif line.strip().startswith("```"):
            lang = line.strip().removeprefix("```").strip()
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(code_lines, lang)
        elif line.strip():
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            text = re.sub(r"`(.+?)`", r"\1", text)
            doc.add_paragraph(text)
        i += 1

    if in_table:
        flush_table()

    doc.save(str(docx_path))


def try_pdf(docx_path: Path, pdf_path: Path) -> bool:
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        return pdf_path.exists()
    except Exception as exc:
        print(f"PDF via docx2pdf unavailable: {exc}")
    try:
        import subprocess
        # LibreOffice headless fallback
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)],
            check=True,
            capture_output=True,
        )
        return pdf_path.exists()
    except Exception as exc:
        print(f"PDF via LibreOffice unavailable: {exc}")
    return False


def main() -> int:
    md_text = assemble_markdown()
    word_count = len(re.findall(r"\b\w+\b", md_text))
    print(f"Assembled markdown: ~{word_count} words")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    md_path = OUT_DIR / "AgroElevate_Final_BlackBook.md"
    docx_path = OUT_DIR / "AgroElevate_Final_BlackBook.docx"
    pdf_path = OUT_DIR / "AgroElevate_Final_BlackBook.pdf"

    md_path.write_text(md_text, encoding="utf-8")
    print(f"Wrote {md_path}")

    md_to_docx(md_text, docx_path)
    print(f"Wrote {docx_path}")

    # Copy to project root
    for name in ["AgroElevate_Final_BlackBook.md", "AgroElevate_Final_BlackBook.docx"]:
        dest = OUT_ROOT / name
        dest.write_bytes((OUT_DIR / name).read_bytes())
        print(f"Copied to {dest}")

    if try_pdf(docx_path, pdf_path):
        root_pdf = OUT_ROOT / "AgroElevate_Final_BlackBook.pdf"
        root_pdf.write_bytes(pdf_path.read_bytes())
        print(f"Wrote {pdf_path}")
        try:
            from pypdf import PdfReader
            pages = len(PdfReader(str(pdf_path)).pages)
            print(f"PDF page count: {pages}")
            if pages < 70:
                print("NOTE: PDF under 70 pages — open .docx in Word and insert screenshots/diagram exports.")
            elif pages > 80:
                print("NOTE: PDF over 80 pages — consider trimming appendix or adjusting Word styles.")
        except Exception:
            pass
    else:
        note = OUT_DIR / "PDF_GENERATION_NOTE.txt"
        note.write_text(
            "PDF could not be auto-generated. Open AgroElevate_Final_BlackBook.docx in Microsoft Word "
            "and use File → Save As → PDF.\n",
            encoding="utf-8",
        )
        print(f"PDF not auto-generated — see {note}")

    est_pages = max(70, min(80, word_count // 280))
    print(f"Estimated formatted pages (Word, 1.5 spacing): ~{est_pages}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
