#!/usr/bin/env python3
"""Edit the original AgroElevate Final Report DOCX in place (technical editor mode)."""
from __future__ import annotations

import re
import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
REPORT_DIR = ROOT / "docs" / "report"
MASTER = REPORT_DIR / "AGROELEVATE_FINAL_REPORT_MASTER.docx"
IMAGES = REPORT_DIR / "images"
OUT_DOCX = ROOT / "AgroElevate_Final_Report_Updated.docx"
OUT_PDF = ROOT / "AgroElevate_Final_Report_Updated.pdf"


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = deepcopy(paragraph._element)
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        new_para.add_run(text)
    return new_para


def insert_heading_after(paragraph: Paragraph, text: str, level: int = 2) -> Paragraph:
    p = insert_paragraph_after(paragraph, text)
    try:
        p.style = f"Heading {level}"
    except Exception:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(14 - level)
    return p


def insert_image_with_caption(
    paragraph: Paragraph, image_path: Path, caption: str, width: float = 5.8
) -> Paragraph:
    img_p = insert_paragraph_after(paragraph, "")
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = insert_paragraph_after(img_p, caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        cap.runs[0].italic = True
    return cap


def insert_table_after(paragraph: Paragraph, headers: list[str], rows: list[list[str]]) -> Paragraph:
    from docx.oxml import OxmlElement

    doc = paragraph.part.document
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for ci, h in enumerate(headers):
        tbl.rows[0].cells[ci].text = h
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            if ci < len(row):
                tbl.rows[ri].cells[ci].text = val
    paragraph._element.addnext(tbl._tbl)
    new_p = OxmlElement("w:p")
    tbl._tbl.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def find_para(doc: Document, pattern: str, occurrence: int = 1, exact: bool = False) -> Paragraph | None:
    count = 0
    pat = pattern.lower()
    for para in doc.paragraphs:
        text = para.text.strip().lower()
        if exact:
            matched = text == pat
        else:
            matched = pat in text
        if matched:
            count += 1
            if count == occurrence:
                return para
    return None


def add_block_after(anchor: Paragraph, blocks: list) -> Paragraph:
    """blocks: list of ('text', str) | ('heading', str, level) | ('image', path, caption) | ('table', headers, rows)"""
    current = anchor
    for block in blocks:
        kind = block[0]
        if kind == "text":
            current = insert_paragraph_after(current, block[1])
        elif kind == "heading":
            level = block[2] if len(block) > 2 else 3
            current = insert_heading_after(current, block[1], level)
        elif kind == "image":
            current = insert_image_with_caption(current, Path(block[1]), block[2], block[3] if len(block) > 3 else 5.8)
        elif kind == "table":
            current = insert_table_after(current, block[1], block[2])
    return current


# ---------------------------------------------------------------------------
# Insertion content (implementation-accurate, additive to original report)
# ---------------------------------------------------------------------------

CH3_WEB_BACKEND = [
    ("heading", "3.6 Web Client Architecture (Implementation)", 3),
    (
        "text",
        "The production web client is implemented as a React 18 Single Page Application using Vite 5 "
        "and TypeScript. The source tree under src/ separates pages (route screens), components "
        "(reusable UI, auth guards, layout shells), hooks (useAuth, useAiService), and lib/ clients "
        "for Supabase, wallet operations, marketplace data, and AI endpoints. React Router 6 provides "
        "client-side navigation; TanStack Query caches server responses with a 60-second staleTime to "
        "reduce redundant database reads. Heavy routes such as Dashboard, Wallet, IntelligenceHub, and "
        "Admin are lazy-loaded with React Suspense, reducing the initial JavaScript bundle by "
        "approximately 69% compared to the eager-loading prototype.",
    ),
    ("heading", "3.7 Backend and Supabase Architecture", 3),
    (
        "text",
        "Supabase hosts PostgreSQL 15 with Row Level Security (RLS), Supabase Auth (JWT sessions), "
        "and Deno Edge Functions. Commerce-critical mutations execute exclusively through SECURITY "
        "DEFINER RPC functions—checkout_order, get_wallet_balance, transfer_funds, and "
        "ensure_profile_from_auth—so wallet balances and royalty splits cannot be altered from the "
        "browser. A dual identity model bridges profiles (UUID, snake_case, RLS anchor) and users "
        "(TEXT uid, camelCase walletBalance) for Android compatibility via ensure_profile_from_auth.",
    ),
    ("image", str(IMAGES / "fig_overall_architecture.jpg"), "Figure 3.5: Updated Overall System Architecture (v1.0.0-rc Implementation)", 4.5),
    ("heading", "3.8 Wallet Architecture", 3),
    (
        "text",
        "Wallet design separates the cached balance on users.walletBalance from the authoritative "
        "append-only ledger wallet_history. Each ledger row records type (deposit, sale, purchase, "
        "royalty_income, transfer, demo_credit), signed amount, optional orderId, and description. "
        "The retired add_funds RPC is blocked for authenticated clients; wallet top-up follows a "
        "server-authoritative Razorpay flow through payment_intents and payment_receipts tables "
        "(Phase G migration 016). Balance updates occur synchronously inside _wallet_ledger_entry "
        "during checkout_order and confirm_wallet_deposit.",
    ),
    ("image", str(IMAGES / "fig_payment_flow.jpg"), "Figure 3.6: Razorpay Wallet Top-Up and Deposit Verification Sequence", 4.5),
    ("heading", "3.9 Royalty Engine Architecture (Option B)", 3),
    (
        "text",
        "AgroElevate implements Option B royalty remittance verified at 12.5% on trader-relisted "
        "inventory purchased by industrialists. When a trader relists produce, products.description "
        "stores JSON metadata including original_farmer_id. During checkout_order, the internal "
        "_commerce_settle_sale function parses this metadata, computes royalty = 12.5% of line total "
        "(clamped 10–12.5% server-side), credits the farmer wallet_history as royalty_income, and "
        "debits the seller net of royalty. Automated verification (npm run commerce:verify) confirms "
        "₹43.75 royalty on a 5 kg × ₹70 resale. Phase 3 adds deferred royalty_obligations linked to "
        "manufacturing_batches for industrialist processing workflows.",
    ),
    ("heading", "3.10 AI Service Architecture", 3),
    (
        "text",
        "The AgroElevate Intelligence API is a FastAPI microservice (v1.0.0-rc) that reads "
        "order_items and products from Supabase, engineers features in feature_engineering.py, runs "
        "scikit-learn models (crop_recommender, demand_intelligence, income_forecaster), and persists "
        "outputs to ai_* tables. Endpoints include /api/intelligence/farmer/dashboard, "
        "/trader/dashboard, /industrialist/dashboard, and POST /api/intelligence/copilot. The React "
        "client wraps calls in withFallback() so commerce remains usable when AI is offline.",
    ),
    ("image", str(IMAGES / "fig_ai_pipeline.jpg"), "Figure 3.7: AI Intelligence Data Pipeline", 4.5),
    ("heading", "3.11 Authentication and Authorization Architecture", 3),
    (
        "text",
        "Registration (Register.tsx) supports five roles: farmer, middleman (trader), industrialist, "
        "customer, and admin. Supabase Auth stores credentials; ensure_profile_from_auth provisions "
        "profiles and users rows. ProtectedRoute enforces session presence; RoleRoute restricts admin "
        "paths. RLS policies scope wallet_history and order_items to the authenticated user or admin "
        "via is_admin(). Suspended and unapproved accounts redirect to /suspended and /pending-approval.",
    ),
    ("image", str(IMAGES / "fig_auth_flow.jpg"), "Figure 3.8: Authentication and Route Guard Flow", 4.5),
    ("heading", "3.12 Payment Gateway Architecture", 3),
    (
        "text",
        "Razorpay integration uses two Supabase Edge Functions: razorpay-create-order (JWT-authenticated, "
        "creates Razorpay order and payment_intents row) and razorpay-webhook (signature-validated, "
        "invokes confirm_wallet_deposit). The web Wallet.tsx opens Razorpay Checkout with the "
        "server-returned order_id only—never a client-constructed order. After payment, the client "
        "polls payment_intents until status = paid and refreshes get_wallet_balance.",
    ),
    ("heading", "3.13 Android Client Architecture (Planned)", 3),
    (
        "text",
        "The Android client (Kotlin, planned for v1.1) shares the same Supabase backend and Razorpay "
        "Edge Function flow documented in ANDROID_RAZORPAY_INTEGRATION.md. Navigation follows Splash → "
        "Login/Register → Role Dashboard → Marketplace, Wallet, Orders, and Intelligence/Copilot. "
        "Offline handling will cache catalog reads locally and queue non-financial actions; all wallet "
        "mutations remain server-side via RPC.",
    ),
    ("image", str(IMAGES / "fig_android_navigation.jpg"), "Figure 3.9: Planned Android Application Navigation Flow", 4.5),
    ("heading", "3.14 Deployment Architecture", 3),
    (
        "text",
        "Development uses Vite dev server (:5173) and Uvicorn AI service (:8000) against Supabase Cloud. "
        "Production deploys the Vite dist/ static bundle to a CDN/static host, Supabase migrations "
        "(001–018) on PostgreSQL, Edge Functions for Razorpay, and the AI service via Docker/Render "
        "(ai-service/Dockerfile, render.yaml). Verification harnesses commerce:verify (26/26 PASS) and "
        "ai:verify validate release readiness.",
    ),
    ("image", str(IMAGES / "fig_deployment.jpg"), "Figure 3.10: Development and Production Deployment Topology", 4.5),
    ("image", str(IMAGES / "fig_marketplace_flow.jpg"), "Figure 3.11: Marketplace Checkout Sequence Diagram", 4.5),
    ("image", str(IMAGES / "fig_order_lifecycle.jpg"), "Figure 3.12: Order Lifecycle State Machine", 4.5),
]

CH3_DATABASE_EXPAND = [
    (
        "text",
        "Implementation database (PostgreSQL 15 on Supabase) uses hybrid naming: camelCase for orders, "
        "order_items, wallet_history, users, crops; snake_case for profiles, products, payment_intents, "
        "payment_receipts, manufacturing_batches, royalty_obligations, and ai_* tables. Primary commerce "
        "RPCs are checkout_order, get_wallet_balance, transfer_funds, ensure_profile_from_auth, "
        "confirm_wallet_deposit, and admin_demo_wallet_credit. No CREATE TRIGGER definitions exist—"
        "business logic is enforced through SECURITY DEFINER functions and CHECK constraints.",
    ),
    ("image", str(IMAGES / "fig_er_diagram.jpg"), "Figure 3.2 (Updated): Entity-Relationship Diagram of Production Schema", 4.5),
    (
        "table",
        ["Table", "Primary Key", "Purpose"],
        [
            ["profiles", "id (UUID)", "Auth-linked identity, role, approval flags"],
            ["users", "uid (TEXT)", "Wallet balance cache, Android bridge"],
            ["products", "id (UUID)", "Marketplace listings, royalty metadata JSON"],
            ["orders", "id (UUID)", "Order header (buyerId, totalAmount, status)"],
            ["order_items", "id (UUID)", "Line items with farmerId, cropId, pricing"],
            ["wallet_history", "id (UUID)", "Append-only wallet ledger"],
            ["payment_intents", "id (UUID)", "Razorpay top-up lifecycle"],
            ["payment_receipts", "id (UUID)", "Confirmed deposit receipts"],
            ["royalty_obligations", "id (UUID)", "Deferred manufacturing royalties"],
            ["manufacturing_batches", "id (UUID)", "Industrialist processing batches"],
            ["ai_crop_recommendations", "id (UUID)", "Persisted AI crop suggestions"],
        ],
    ),
]

CH4_TECH_ALIGNMENT = [
    (
        "text",
        "Implementation alignment note (v1.0.0-rc): The deployed web application extends the technologies "
        "listed above with React 18, Vite 5, TypeScript 5.8, TanStack Query 5, shadcn/ui components, "
        "and a FastAPI AI microservice. Tailwind CSS remains the styling layer; HTML/CSS/JavaScript "
        "concepts described earlier map to the React component architecture in src/pages and src/components.",
    ),
]

CH4_MODULES = [
    ("heading", "4.7 Farmer Module Implementation", 3),
    (
        "text",
        "Purpose: Enable cultivators to list produce, receive sale proceeds, and collect downstream "
        "royalties automatically. Workflow: Register with role farmer → ensure_profile_from_auth → "
        "INSERT products → buyer checkout_order credits seller wallet_history type sale. When traders "
        "relist with metadata, farmer receives royalty_income on industrialist purchase. Screens: "
        "/dashboard (FarmerDashboardSection), /marketplace (My Listings tab), /wallet, /intelligence. "
        "Validation: bank account required at registration; RLS restricts listings to seller_id = auth.uid().",
    ),
    ("text", "[Screenshot placeholder: Figure 4.5 — Farmer Dashboard with sales KPIs and royalty summary]"),
    ("heading", "4.8 Trader (Middleman) Module Implementation", 3),
    (
        "text",
        "Purpose: Procure from farmers, manage inventory, relist with ownership metadata for royalty "
        "enforcement. Role stored as middleman in profiles and trader in users. Workflow: Razorpay "
        "wallet top-up → purchase via checkout_order → relist with products.description JSON containing "
        "original_farmer_id, source_order_id, royalty_rate (12.5%). Screens: /dashboard, /marketplace, "
        "/wallet, TraderInsights on /intelligence. Error handling: insufficient balance prompts wallet "
        "top-up before checkout proceeds.",
    ),
    ("text", "[Screenshot placeholder: Figure 4.6 — Trader inventory relist dialog with royalty metadata]"),
    ("heading", "4.9 Industrialist Module Implementation", 3),
    (
        "text",
        "Purpose: Procure agricultural input, run manufacturing batches, sell processed goods with "
        "deferred royalty settlement. Workflow: marketplace purchase triggers immediate 12.5% farmer "
        "royalty on relisted trader goods; direct farmer purchase creates royalty_obligations for "
        "manufacturing_batches. Screens: /dashboard, /marketplace, IndustrialistInsights. Database: "
        "manufacturing_batches, processed_products, royalty_obligations (migration 014).",
    ),
    ("text", "[Screenshot placeholder: Figure 4.7 — Industrialist manufacturing and procurement panel]"),
    ("heading", "4.10 Customer Module Implementation", 3),
    (
        "text",
        "Purpose: Retail/end-user commerce without intelligence or relist privileges. Added in Option B "
        "Phase 1 (migration 012). Workflow: Register as customer → browse /marketplace → cart checkout "
        "via checkout_order → view /orders. Wallet top-up uses the same Razorpay flow as business roles. "
        "Screens: /dashboard (customer variant), /marketplace, /wallet, /orders.",
    ),
    ("text", "[Screenshot placeholder: Figure 4.8 — Customer marketplace browse and checkout]"),
    ("heading", "4.11 Admin Dashboard Implementation", 3),
    (
        "text",
        "Purpose: User moderation, payment audit, and demo wallet credits for academic demonstrations. "
        "Routes: /admin (profile suspend/approve), /admin/payments (payment_intents and receipts audit, "
        "demo credit via admin_demo_wallet_credit RPC, migrations 017–018). Authorization: RoleRoute "
        "allowedRole=admin plus is_admin() RLS policies.",
    ),
    ("text", "[Screenshot placeholder: Figure 4.9 — Admin user moderation and payments console]"),
    ("heading", "4.12 Wallet Module Implementation", 3),
    (
        "text",
        "Purpose: Server-authoritative stored-value wallet for all commerce. APIs: get_wallet_balance, "
        "transfer_funds, Razorpay Edge Functions, confirm_wallet_deposit (webhook). UI: Wallet.tsx "
        "shows balance, top-up modal, transaction history with type badges (deposit, royalty_income, "
        "demo_credit). Security: direct add_funds blocked; verified in commerce harness.",
    ),
    ("text", "[Screenshot placeholder: Figure 4.10 — Wallet balance, Razorpay top-up, and transaction ledger]"),
    ("heading", "4.13 Royalty Engine Implementation", 3),
    (
        "text",
        "Algorithm (checkout_order → _commerce_settle_sale): For each line item, parse "
        "_parse_product_commerce_meta from products.description; if original_farmer_id present and "
        "buyer role qualifies, royalty_amount = ROUND(line_total × 0.125, 2); transfer royalty via "
        "_wallet_transfer to farmer; credit seller net = line_total − royalty. Verified: 5×₹70 = ₹350 "
        "→ ₹43.75 royalty_income. Harness: commerce:verify test 26/26 PASS.",
    ),
    ("heading", "4.14 AI and Copilot Implementation", 3),
    (
        "text",
        "FastAPI intelligence_service aggregates order history, runs feature_engineering, executes "
        "sklearn models, and returns role dashboards. Copilot (POST /api/intelligence/copilot) provides "
        "grounded conversational assistance using the same feature context. Insufficient-data states "
        "display when training rows are below threshold. Persistence: ai_crop_recommendations, "
        "ai_market_predictions, ai_income_forecasts with user-scoped RLS.",
    ),
    ("text", "[Screenshot placeholder: Figure 4.11 — Farmer Intelligence dashboard and crop recommendations]"),
    ("text", "[Screenshot placeholder: Figure 4.12 — AI Copilot chat panel]"),
    ("heading", "4.15 Android Application Implementation", 3),
    (
        "text",
        "The Android application is architecturally specified to mirror the web client's Supabase and "
        "Razorpay integration (ANDROID_RAZORPAY_INTEGRATION.md). Planned folder structure: app/src/main "
        "with activities/fragments for Login, Dashboard, Marketplace, Wallet, Orders, Notifications; "
        "data layer using Supabase Kotlin SDK; Razorpay Standard SDK for checkout. REST communication "
        "uses the same RPC and Edge Function endpoints. Offline: read-only catalog cache; financial "
        "operations require connectivity. Testing: instrumented UI tests for auth and marketplace browse; "
        "wallet tests against Supabase test project.",
    ),
    ("text", "[Screenshot placeholder: Figure 4.13 — Android login screen (planned)]"),
    ("text", "[Screenshot placeholder: Figure 4.14 — Android marketplace screen (planned)]"),
    ("heading", "4.16 Authentication and Authorization Implementation", 3),
    (
        "text",
        "Login algorithm: signInWithPassword → ensure_profile_from_auth → check profiles.suspended and "
        "profiles.approved → route to role dashboard. Registration algorithm: signUp with metadata → "
        "ensureUserRecords inserts profiles + users → email verification path when session null. "
        "Authorization layers: JWT (Supabase), RLS (PostgreSQL), RPC SECURITY DEFINER (commerce), "
        "client guards (ProtectedRoute, RoleRoute).",
    ),
    ("heading", "4.17 Notifications and Analytics", 3),
    (
        "text",
        "Notifications: In-app toast feedback on checkout success and royalty credit; wallet_history "
        "serves as durable financial notification. Analytics: Role dashboards aggregate order_items and "
        "wallet_history; IntelligenceHub provides AI-driven charts (ThemedChart.tsx, Recharts). Admin "
        "payment audit exports payment_receipts for reconciliation.",
    ),
    ("heading", "4.18 Testing and Verification", 3),
    (
        "text",
        "Automated verification is implemented as Node.js harness scripts rather than unit-test frameworks "
        "alone. npm run commerce:verify executes 26 end-to-end checks including registration, wallet "
        "top-up simulation, checkout, royalty math, RLS isolation, and retired add_funds guard. "
        "npm run commerce:smoke validates seven RPC contracts. npm run ai:verify checks /health and "
        "farmer dashboard. npm run build confirms production TypeScript compilation.",
    ),
    (
        "table",
        ["Test ID", "Category", "Description", "Expected", "Result"],
        [
            ["CV-01", "Integration", "Farmer registration + wallet provision", "profiles + users row", "PASS"],
            ["CV-12", "Royalty", "Trader relist → industrialist buy 5×₹70", "₹43.75 royalty_income", "PASS"],
            ["CV-18", "Security", "Authenticated add_funds RPC", "Exception raised", "PASS"],
            ["CV-26", "System", "Full commerce regression suite", "26/26 checks", "PASS"],
            ["SM-01", "Smoke", "checkout_order RPC contract", "Valid JSON response", "PASS"],
            ["AI-01", "AI", "GET /health + farmer dashboard", "200 OK", "PASS"],
            ["BLD-01", "Build", "npm run build", "Zero TS errors", "PASS"],
        ],
    ),
    ("heading", "4.19 Security and Performance Testing", 3),
    (
        "text",
        "Security testing: cross-user wallet_history reads blocked by RLS; service role key absent from "
        "Vite dist bundle; Razorpay webhook requires valid signature. Performance testing: lazy routes "
        "reduced main chunk from ~1,256 KB to ~384 KB; TanStack Query 60s staleTime lowers Supabase "
        "round trips. User acceptance: manual commerce scenarios documented in MANUAL_COMMERCE_TEST.md.",
    ),
]

CH5_RESULTS_EXPAND = [
    ("heading", "5.4 Wallet and Payment Results", 3),
    (
        "text",
        "Wallet operations were validated through automated harness and manual Razorpay Test Mode "
        "top-ups. get_wallet_balance reconciles with wallet_history sums. payment_intents lifecycle "
        "(created → paid) completes via webhook or polling. Demo credits via admin_demo_wallet_credit "
        "produce auditable demo_credit ledger rows (migrations 017–018).",
    ),
    ("text", "[Screenshot placeholder: Figure 5.4 — Wallet transaction history showing deposit and purchase entries]"),
    ("heading", "5.5 Royalty Redistribution Results", 3),
    (
        "text",
        "Royalty verification confirms Option B Rule 3: industrialist purchase of trader-relisted "
        "inventory yields 12.5% immediate credit to the original farmer. Measured result: ₹43.75 on "
        "₹350 line total (5 kg × ₹70/kg), within ±₹0.02 harness tolerance. Farmer wallet_history "
        "type royalty_income appears without manual intervention.",
    ),
    ("heading", "5.6 AI Intelligence Results", 3),
    (
        "text",
        "AI service v1.0.0-rc returns crop recommendations, demand signals, and income forecasts when "
        "sufficient order history exists; otherwise dashboards display explicit insufficient-data states. "
        "npm run ai:verify confirms /health and farmer dashboard endpoints. Copilot responds with "
        "role-aware context drawn from the same feature pipeline.",
    ),
    ("text", "[Screenshot placeholder: Figure 5.5 — AI crop recommendation panel with confidence scores]"),
    ("heading", "5.7 Marketplace and Multi-Role Commerce Results", 3),
    (
        "text",
        "Marketplace checkout completes atomically across farmer, trader, industrialist, and customer "
        "roles. products.quantity decrements in the same transaction as wallet debits. Cart checkout "
        "redirects to /orders with visible order_items. Customer role (migration 012) participates "
        "in the same checkout_order RPC as business roles.",
    ),
    (
        "table",
        ["Verification Harness", "Tests", "Outcome", "Release"],
        [
            ["commerce:verify", "26", "26/26 PASS", "v1.0.0-rc"],
            ["commerce:smoke", "7", "7/7 PASS", "v1.0.0-rc"],
            ["ai:verify", "Health + dashboard", "PASS", "v1.0.0-rc"],
            ["npm run build", "Production compile", "PASS", "v1.0.0-rc"],
        ],
    ),
    ("heading", "5.8 Android and Deployment Readiness", 3),
    (
        "text",
        "Android client architecture and Razorpay integration are documented and API-compatible; "
        "native source delivery is planned for v1.1. Web platform readiness score: 86/100 production, "
        "90/100 demo (FINAL_RELEASE_REPORT.md). Deployment artifacts include Vite dist/, Supabase "
        "migrations 001–018, Edge Functions, and Dockerized AI service.",
    ),
]


def apply_insertions(doc: Document) -> int:
    """Return count of insertion blocks applied."""
    applied = 0
    plan = [
        ("3.3 database architecture", 1, True, CH3_DATABASE_EXPAND),
        ("3.5 system workflow", 1, True, [
            ("text", "The following diagrams document the implemented checkout and order settlement flows verified in commerce:verify."),
        ]),
        ("the workflow concludes when all records have been successfully updated", 1, False, CH3_WEB_BACKEND),
        ("the integration of html, css, tailwind css, javascript, supabase, vs code, and google chrome provided", 1, False, CH4_TECH_ALIGNMENT),
        ("the implementation of this module demonstrates the practical application of digital commerce principles within the agricultural sector", 1, False, CH4_MODULES),
        ("5.3 impact analysis of royalty redistribution model", 1, True, CH5_RESULTS_EXPAND),
    ]

    anchors: list[tuple[Paragraph, list]] = []
    for pattern, occ, exact, blocks in plan:
        para = find_para(doc, pattern, occ, exact=exact)
        if para is None:
            print(f"WARNING: anchor not found: {pattern!r}")
            continue
        anchors.append((para, blocks))

    # Fallback: marketplace module closing paragraph (Ch. 4)
    if not any("4.7 Farmer Module" in str(b) for _, b in anchors for b in b if b[0] == "heading"):
        if len(doc.paragraphs) > 324:
            anchors.append((doc.paragraphs[324], CH4_MODULES))
            print("Fallback anchor: paragraph 324 (Marketplace module)")

    for para, blocks in anchors:
        add_block_after(para, blocks)
        applied += 1
        print(f"Inserted after: {para.text[:70]!r}...")

    return applied


def update_abstract_addendum(doc: Document) -> None:
    para = find_para(
        doc,
        "agroelevate is an ai-powered agri-commerce platform for farmer transactions and analytics",
        1,
    )
    if not para:
        para = find_para(doc, "in addition to providing marketplace functionality, agroelevate analyzes", 1)
    if para:
        insert_paragraph_after(
            para,
            "The v1.0.0-rc implementation additionally delivers a React/Vite/TypeScript web client, "
            "Razorpay wallet top-ups, Option B royalty at 12.5% (verified ₹43.75 on the standard "
            "harness case), FastAPI intelligence services with Copilot, and automated commerce "
            "verification (26/26 PASS).",
        )


def main() -> int:
    if not MASTER.exists():
        print(f"Master DOCX not found: {MASTER}")
        print("Place AGROELEVATE_FINAL_REPORT.pdf in Downloads and re-run conversion.")
        return 1

    shutil.copy2(MASTER, REPORT_DIR / "AGROELEVATE_FINAL_REPORT_WORKING.docx")
    doc = Document(str(REPORT_DIR / "AGROELEVATE_FINAL_REPORT_WORKING.docx"))
    print(f"Loaded master report: {len(doc.paragraphs)} paragraphs")

    update_abstract_addendum(doc)
    n = apply_insertions(doc)

    doc.save(str(OUT_DOCX))
    print(f"Saved {OUT_DOCX} ({OUT_DOCX.stat().st_size} bytes), {n} insertion blocks")

    try:
        from docx2pdf import convert
        from pypdf import PdfReader

        convert(str(OUT_DOCX), str(OUT_PDF))
        pages = len(PdfReader(str(OUT_PDF)).pages)
        print(f"Saved {OUT_PDF} — {pages} pages")
    except Exception as exc:
        print(f"PDF generation note: {exc}")
        note = REPORT_DIR / "PDF_GENERATION_NOTE.txt"
        note.write_text(
            "Open AgroElevate_Final_Report_Updated.docx in Word → Save As PDF.\n",
            encoding="utf-8",
        )

    return 0


if __name__ == "__main__":
    sys.exit(main())
